import json
import re
from docx import Document

# -----------------------------------------
# STEP 1: Load Document
# -----------------------------------------

def load_doc(path):
    print("\n🟢 Loading document...")
    doc = Document(path)
    print("Total paragraphs:", len(doc.paragraphs))
    print("Total tables:", len(doc.tables))
    return doc


# -----------------------------------------
# STEP 2: Extract text from paragraphs
# -----------------------------------------

def extract_paragraph_text(doc):
    print("\n🟡 Extracting paragraphs...")
    text_list = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            text_list.append(t)

    print("Paragraphs extracted:", len(text_list))
    return text_list


# -----------------------------------------
# STEP 3: Extract text from tables
# -----------------------------------------

def extract_table_text(doc):
    print("\n🟠 Extracting tables...")
    text_list = []

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    text_list.append(t)

    print("Table cells extracted:", len(text_list))
    return text_list


# -----------------------------------------
# STEP 4: Combine both sources
# -----------------------------------------

def combine_text(paragraphs, table_text):
    combined = paragraphs + table_text
    print("\n🔵 Total text blocks combined:", len(combined))
    return combined


# -----------------------------------------
# STEP 5: Cleaning Helpers
# -----------------------------------------

REQ_PATTERN = r'\b(FR|NFR|SR|DR|IR)-\d+\b'

def clean_requirement_text(text):
    # Remove leading dots, bullets, symbols
    text = re.sub(r'^[\.\-\•\:\s]+', '', text)

    # Remove table header residue like "ID Requirement"
    text = re.sub(r'\bID Requirement\b', '', text, flags=re.IGNORECASE)

    # Remove multiple spaces
    text = re.sub(r'\s+', ' ', text)

    return text.strip()


def is_valid_requirement(text):
    # Remove empty or very short garbage
    if not text:
        return False

    if len(text.split()) < 5:
        return False

    # Remove pure header words
    if text.lower() in ["requirement", "id", "id requirement"]:
        return False

    return True


# -----------------------------------------
# STEP 6: Extract requirements (Improved)
# -----------------------------------------

def extract_requirements(text_blocks):
    print("\n🟣 Extracting requirements (Improved)...")

    req_dict = {}
    current_id = None

    for line in text_blocks:
        match = re.search(REQ_PATTERN, line)

        if match:
            req_id = match.group()
            text_after = line.split(req_id, 1)[-1]

            cleaned_text = clean_requirement_text(text_after)

            req_dict[req_id] = cleaned_text
            current_id = req_id

        else:
            if current_id:
                req_dict[current_id] += " " + line

    print("Total Raw Extracted:", len(req_dict))

    # Final Cleaning + Validation
    final_dict = {}
    removed_count = 0

    for req_id, text in req_dict.items():
        cleaned_text = clean_requirement_text(text)

        if is_valid_requirement(cleaned_text):
            final_dict[req_id] = cleaned_text
        else:
            removed_count += 1

    print("Valid Requirements:", len(final_dict))
    print("Removed Invalid:", removed_count)

    return final_dict


# -----------------------------------------
# STEP 7: Save to JSON
# -----------------------------------------

def save_json(data, path="requirements.json"):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=4, ensure_ascii=False)
    print("\n🟤 Saved output to:", path)


# -----------------------------------------
# STEP 8: Print samples
# -----------------------------------------

def print_samples(req_dict, n=10):
    print("\n🔴 SAMPLE REQUIREMENTS:")
    print("-------------------------")
    for i, (k, v) in enumerate(req_dict.items()):
        print(f"{k} → {v[:120]}...")
        if i >= n - 1:
            break


# -----------------------------------------
# MAIN PIPELINE
# -----------------------------------------

if __name__ == "__main__":
    doc = load_doc("SRS.docx")
    paragraphs = extract_paragraph_text(doc)
    table_text = extract_table_text(doc)
    all_blocks = combine_text(paragraphs, table_text)
    requirements = extract_requirements(all_blocks)
    save_json(requirements)
    print_samples(requirements)