import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("🟢 Loading files...")

# =========================
# LOAD FILES SAFELY
# =========================
try:
    with open("requirements.json", "r", encoding="utf-8") as f:
        requirements = json.load(f)
    print("✅ requirements.json loaded")

    with open("duplicate_report.json", "r", encoding="utf-8") as f:
        duplicates_grouped = json.load(f)
    print("✅ duplicate_report.json loaded")

    with open("ambiguity_report.json", "r", encoding="utf-8") as f:
        ambiguity_data = json.load(f)
    print("✅ ambiguity_report.json loaded")

except Exception as e:
    print("❌ Error loading files:", e)
    exit()

# =========================
# NORMALIZE REQUIREMENTS
# =========================
if isinstance(requirements, list):
    requirements = {item["id"]: item["text"] for item in requirements}

print(f"📌 Total raw requirements: {len(requirements)}")

# =========================
# PROCESS DUPLICATES
# =========================
duplicate_pairs = []
duplicate_removed_ids = set()

# Handle duplicate report format: dict with 'duplicates' key or list of groups
if isinstance(duplicates_grouped, dict) and "duplicates" in duplicates_grouped:
    duplicate_groups = duplicates_grouped["duplicates"]
elif isinstance(duplicates_grouped, list):
    duplicate_groups = duplicates_grouped
else:
    print("❌ Unknown duplicate report format")
    duplicate_groups = []

for group in duplicate_groups:
    if not isinstance(group, list) or len(group) < 2:
        continue

    # Keep first requirement as master, remove others
    master_req = group[0]["id"]
    for dup_req in group[1:]:
        dup_id = dup_req.get("id")
        if dup_id:
            duplicate_pairs.append({"req1": master_req, "req2": dup_id})
            duplicate_removed_ids.add(dup_id)

print(f"🧹 Duplicate IDs marked: {len(duplicate_removed_ids)}")

# =========================
# PROCESS AMBIGUOUS REQUIREMENTS
# =========================
ambiguous_requirements = []

for item in ambiguity_data:
    req_id = item.get("id")
    if not req_id or req_id in duplicate_removed_ids:
        continue

    ambiguous_requirements.append({
        "id": req_id,
        "text": item.get("text", ""),
        "ambiguous_flags": item.get("ambiguous_flags", []),
        "ambiguity_score": item.get("ambiguity_score", 0)
    })

# Sort ambiguous by score descending
ambiguous_requirements.sort(key=lambda x: x["ambiguity_score"], reverse=True)

# =========================
# FORMAT VALIDATION
# =========================
format_issues = []

for req_id, text in requirements.items():
    if req_id in duplicate_removed_ids:
        continue

    text_lower = text.lower()
    if "shall" not in text_lower and "should" not in text_lower:
        format_issues.append({
            "id": req_id,
            "text": text,
            "reason": "Missing SHALL/SHOULD keyword"
        })

# =========================
# BUILD REPORT DICTIONARY
# =========================
annotated_srs = {
    "summary": {
        "total_requirements": len(requirements),
        "duplicate_groups": len(duplicate_groups),
        "duplicates_removed": len(duplicate_removed_ids),
        "ambiguous_count": len(ambiguous_requirements),
        "format_issue_count": len(format_issues)
    },
    "duplicates": duplicate_pairs,
    "ambiguous_requirements": ambiguous_requirements,
    "format_issues": format_issues
}

# =========================
# SAVE JSON
# =========================
with open("annotated_srs.json", "w", encoding="utf-8") as f:
    json.dump(annotated_srs, f, indent=4, ensure_ascii=False)

print("✅ annotated_srs.json saved")

# =========================
# GENERATE DOCX
# =========================
print("📝 Generating annotated_srs.docx...")

doc = Document()
doc.add_heading("Annotated Software Requirements Specification", level=1)

# Summary
doc.add_heading("1. Summary", level=2)
doc.add_paragraph(f"Total Requirements: {annotated_srs['summary']['total_requirements']}")
doc.add_paragraph(f"Duplicate Groups: {annotated_srs['summary']['duplicate_groups']}")
doc.add_paragraph(f"Duplicates Removed: {annotated_srs['summary']['duplicates_removed']}")
doc.add_paragraph(f"Ambiguous Requirements: {annotated_srs['summary']['ambiguous_count']}")
doc.add_paragraph(f"Format Issues: {annotated_srs['summary']['format_issue_count']}")

# Duplicates
doc.add_heading("2. Duplicate Requirements", level=2)
if duplicate_pairs:
    for pair in duplicate_pairs:
        doc.add_paragraph(f"{pair['req1']} is duplicate of {pair['req2']}", style='List Bullet')
else:
    doc.add_paragraph("No duplicates found.")

# Ambiguous
doc.add_heading("3. Ambiguous Requirements", level=2)
if ambiguous_requirements:
    for item in ambiguous_requirements:
        doc.add_heading(item["id"], level=3)
        doc.add_paragraph(f"Requirement: {item['text']}")
        doc.add_paragraph(f"Ambiguity Score: {item['ambiguity_score']}")
        if item["ambiguous_flags"]:
            doc.add_paragraph("Ambiguity Flags:")
            for flag in item["ambiguous_flags"]:
                doc.add_paragraph(flag, style='List Bullet')
else:
    doc.add_paragraph("No ambiguous requirements found.")

# Format Issues
doc.add_heading("4. Format Issues", level=2)
if format_issues:
    for item in format_issues:
        doc.add_heading(item["id"], level=3)
        doc.add_paragraph(f"Requirement: {item['text']}")
        doc.add_paragraph(f"Issue: {item['reason']}")
else:
    doc.add_paragraph("No format issues found.")

doc.save("annotated_srs.docx")
print("✅ annotated_srs.docx generated successfully!")
print("\n🎉 DONE!")